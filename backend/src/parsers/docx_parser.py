"""Word (.docx) parser using python-docx heading styles."""

import io
import logging
from pathlib import PurePosixPath
from typing import Awaitable, Callable

from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

from src.parsers.base import (
    DocumentAsset,
    DocumentBlock,
    DocumentSection,
    ParsedDocument,
    blocks_to_sections,
)

logger = logging.getLogger(__name__)

# Mapping from python-docx style names to heading levels
HEADING_STYLE_MAP = {
    "Heading 1": 1,
    "Heading 2": 2,
    "Heading 3": 3,
    "Heading 4": 3,
    "Title": 1,
}


class DocxParser:
    """Extract structured sections from Word documents."""

    def parse(self, content: bytes) -> list[DocumentSection]:
        import asyncio

        try:
            asyncio.get_running_loop()
        except RuntimeError:
            return blocks_to_sections(asyncio.run(self.parse_blocks(content)))
        raise RuntimeError(
            "DocxParser.parse() cannot be called from a running event loop; "
            "use parse_blocks() instead."
        ) from None

    async def parse_blocks(
        self,
        content: bytes,
        asset_saver: Callable[[str, bytes], Awaitable[str]] | None = None,
        filename: str = "document.docx",
    ) -> ParsedDocument:
        """Parse DOCX into reading-order blocks."""
        from docx import Document

        doc = Document(io.BytesIO(content))
        logger.info("[DocxParser] start parse_blocks: filename=%s", filename)
        blocks: list[DocumentBlock] = []
        assets: list[DocumentAsset] = []
        active_title = ""
        order = 0
        table_count = 0
        image_count = 0

        for element in doc.element.body.iterchildren():
            tag = element.tag
            if tag == qn("w:p"):
                para = Paragraph(element, doc)
                paragraph_blocks, extracted_images = await self._parse_paragraph(
                    para=para,
                    doc=doc,
                    active_title=active_title,
                    order=order,
                    asset_saver=asset_saver,
                )
                for block in paragraph_blocks:
                    blocks.append(block)
                    if block.type in ("image", "chart") and block.asset_path:
                        assets.append(
                            DocumentAsset(
                                id=block.id,
                                type=block.type,
                                path=block.asset_path,
                                filename=PurePosixPath(block.asset_path).name,
                            )
                        )
                    order += 1
                    if block.type == "title":
                        active_title = block.text
                    logger.info(
                        "[DocxParser] block parsed: order=%d type=%s level=%d text=%s",
                        block.order,
                        block.type,
                        block.level,
                        block.text[:80],
                    )
                image_count += extracted_images
            elif tag == qn("w:tbl"):
                table = Table(element, doc)
                block = self._table_to_block(table, order, active_title)
                blocks.append(block)
                table_count += 1
                logger.info(
                    "[DocxParser] table parsed: order=%d rows=%d cols=%d parent=%s",
                    block.order,
                    len(table.rows),
                    len(table.columns),
                    active_title,
                )
                order += 1

        parsed = ParsedDocument(title=filename, blocks=blocks, assets=assets)
        logger.info(
            "[DocxParser] done: blocks=%d tables=%d images=%d",
            len(blocks),
            table_count,
            image_count,
        )
        return parsed

    async def _parse_paragraph(
        self,
        para: Paragraph,
        doc,
        active_title: str,
        order: int,
        asset_saver: Callable[[str, bytes], Awaitable[str]] | None,
    ) -> tuple[list[DocumentBlock], int]:
        blocks: list[DocumentBlock] = []
        image_count = 0
        text = para.text.strip()
        style_name = para.style.name if para.style else ""
        level = HEADING_STYLE_MAP.get(style_name)

        if text:
            block_type = "title" if level is not None else self._paragraph_type(style_name)
            blocks.append(
                DocumentBlock(
                    id=f"docx-{order}",
                    type=block_type,
                    text=text,
                    level=level or 0,
                    parent_title="" if block_type == "title" else active_title,
                    order=order,
                )
            )
            order += 1

        for rel_id in self._image_rel_ids(para):
            image_part = doc.part.related_parts.get(rel_id)
            if image_part is None:
                logger.warning("[DocxParser] image rel missing: rel_id=%s", rel_id)
                continue
            part_name = str(getattr(image_part, "partname", "image.png"))
            suffix = PurePosixPath(part_name).suffix or ".png"
            image_count += 1
            image_filename = f"image_{order:04d}{suffix}"
            asset_path = ""
            if asset_saver:
                asset_path = await asset_saver(image_filename, image_part.blob)
            block = DocumentBlock(
                id=f"docx-{order}",
                type="image",
                caption=text if text and len(text) <= 120 else "",
                asset_path=asset_path,
                parent_title=active_title,
                order=order,
            )
            blocks.append(block)
            logger.info(
                "[DocxParser] image extracted: order=%d rel_id=%s size=%d path=%s",
                block.order,
                rel_id,
                len(image_part.blob),
                asset_path,
            )
            order += 1

        return blocks, image_count

    def _paragraph_type(self, style_name: str) -> str:
        lowered = style_name.lower()
        if "list" in lowered or "bullet" in lowered:
            return "list"
        return "paragraph"

    def _image_rel_ids(self, para: Paragraph) -> list[str]:
        rel_ids: list[str] = []
        for blip in para._element.xpath(".//*[local-name()='blip']"):
            rel_id = blip.get(qn("r:embed")) or blip.get(qn("r:id"))
            if rel_id and rel_id not in rel_ids:
                rel_ids.append(rel_id)
        for image_data in para._element.xpath(".//*[local-name()='imagedata']"):
            rel_id = image_data.get(qn("r:id"))
            if rel_id and rel_id not in rel_ids:
                rel_ids.append(rel_id)
        return rel_ids

    def _table_to_block(self, table: Table, order: int, active_title: str) -> DocumentBlock:
        rows = [[cell.text.strip().replace("\n", " ") for cell in row.cells] for row in table.rows]
        markdown = self._table_to_markdown(rows)
        html = self._table_to_html(rows)
        summary = self._summarize_table(rows)
        return DocumentBlock(
            id=f"docx-{order}",
            type="table",
            text=markdown,
            parent_title=active_title,
            html=html,
            summary=summary,
            order=order,
        )

    def _table_to_markdown(self, rows: list[list[str]]) -> str:
        if not rows:
            return ""
        width = max(len(row) for row in rows)
        normalized = [row + [""] * (width - len(row)) for row in rows]
        header = "| " + " | ".join(normalized[0]) + " |"
        sep = "| " + " | ".join("---" for _ in range(width)) + " |"
        body = ["| " + " | ".join(row) + " |" for row in normalized[1:]]
        return "\n".join([header, sep, *body])

    def _table_to_html(self, rows: list[list[str]]) -> str:
        html_rows = []
        for row in rows:
            cells = "".join(f"<td>{cell}</td>" for cell in row)
            html_rows.append(f"<tr>{cells}</tr>")
        return "<table>" + "".join(html_rows) + "</table>"

    def _summarize_table(self, rows: list[list[str]]) -> str:
        if not rows:
            return ""
        col_count = max(len(row) for row in rows)
        headers = [cell for cell in rows[0] if cell]
        header_text = "、".join(headers[:8])
        return f"表格共 {len(rows)} 行 {col_count} 列，字段包括：{header_text}。" if header_text else f"表格共 {len(rows)} 行 {col_count} 列。"

    def _parse_legacy(self, content: bytes) -> list[DocumentSection]:
        from docx import Document

        doc = Document(io.BytesIO(content))
        sections: list[DocumentSection] = []
        current_title = ""
        current_level = 1
        current_lines: list[str] = []
        chapter_title = ""

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""
            level = HEADING_STYLE_MAP.get(style_name)

            if level is not None:
                # Flush previous section
                if current_lines:
                    sections.append(
                        DocumentSection(
                            title=current_title,
                            level=current_level,
                            content="\n".join(current_lines),
                            parent_title=chapter_title if current_level > 1 else "",
                        )
                    )
                    current_lines = []

                current_title = text
                current_level = level
                if level == 1:
                    chapter_title = text
            else:
                current_lines.append(text)

        # Flush last section
        if current_lines:
            sections.append(
                DocumentSection(
                    title=current_title,
                    level=current_level,
                    content="\n".join(current_lines),
                    parent_title=chapter_title if current_level > 1 else "",
                )
            )

        # Fallback: if no headings detected, return entire doc as one section
        if not sections:
            all_text = "\n".join(
                p.text.strip() for p in doc.paragraphs if p.text.strip()
            )
            if all_text:
                sections.append(
                    DocumentSection(title="", level=1, content=all_text)
                )

        logger.info("[DocxParser] extracted %d sections", len(sections))
        return sections
