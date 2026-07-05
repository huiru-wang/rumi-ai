import io

import pytest
from docx import Document

from src.parsers.docx_parser import DocxParser


@pytest.mark.asyncio
async def test_docx_parser_extracts_heading_paragraph_and_table_blocks():
    doc = Document()
    doc.add_heading("第一章 概述", level=1)
    doc.add_paragraph("这是正文段落。")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "指标"
    table.cell(0, 1).text = "数值"
    table.cell(1, 0).text = "A"
    table.cell(1, 1).text = "10"
    buffer = io.BytesIO()
    doc.save(buffer)

    parsed = await DocxParser().parse_blocks(buffer.getvalue())

    assert [block.type for block in parsed.blocks] == ["title", "paragraph", "table"]
    assert parsed.blocks[0].level == 1
    assert parsed.blocks[1].parent_title == "第一章 概述"
    assert "| 指标 | 数值 |" in parsed.blocks[2].text
    assert "<table>" in parsed.blocks[2].html
